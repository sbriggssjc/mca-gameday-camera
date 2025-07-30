import argparse
import json
import os
from pathlib import Path
from typing import List

try:
    from llama_index import (
        VectorStoreIndex,
        SimpleDirectoryReader,
        ServiceContext,
        Document
    )
except Exception:  # pragma: no cover - optional dependency
    VectorStoreIndex = None
    SimpleDirectoryReader = None
    ServiceContext = None
    Document = None

try:
    import openai
except Exception:  # pragma: no cover - optional dependency
    openai = None

try:
    import pyttsx3  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    pyttsx3 = None

try:
    import vosk  # type: ignore
    import pyaudio  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    vosk = None
    pyaudio = None


DEFAULT_SOURCES = [
    "playbook.docx",
    "outputs/play_metadata.json",
    "outputs/predictions.json",
]


def _load_documents(paths: List[str]) -> List[Document]:
    """Load text documents from the provided paths."""
    docs: List[Document] = []
    for path in paths:
        if not os.path.exists(path):
            continue
        if path.endswith(('.docx', '.pptx')):
            try:
                from docx import Document as Docx

                doc_obj = Docx(path)
                text = "\n".join(p.text for p in doc_obj.paragraphs)
                docs.append(Document(text=text, metadata={"source": path}))
            except Exception as exc:  # pragma: no cover - optional dependency
                print(f"[!] Failed to read {path}: {exc}")
        else:
            try:
                with open(path, "r", encoding="utf-8") as f:
                    docs.append(Document(text=f.read(), metadata={"source": path}))
            except Exception as exc:  # pragma: no cover - optional dependency
                print(f"[!] Failed to read {path}: {exc}")
    return docs


def answer_question(query: str, context_sources: List[str] | None = None) -> str:
    """Answer a question using local documents and GPT-4."""
    if VectorStoreIndex is None:
        raise ImportError("llama_index is required for answer_question")

    sources = list(DEFAULT_SOURCES)
    if context_sources:
        sources.extend(context_sources)

    documents = _load_documents(sources)
    if not documents:
        return "I have no context to answer that question."

    service_context = ServiceContext.from_defaults()
    index = VectorStoreIndex.from_documents(documents, service_context=service_context)
    query_engine = index.as_query_engine()
    response = query_engine.query(query)
    return str(response)


def _speak(text: str) -> None:
    if pyttsx3 is None:
        print(text)
        return
    engine = pyttsx3.init()
    engine.say(text)
    engine.runAndWait()


def listen_and_respond() -> None:
    """Listen on microphone and respond via TTS."""
    if vosk is None or pyaudio is None:
        print("[!] Vosk and PyAudio are required for live mode")
        return

    model = vosk.Model(lang="en-us")
    recognizer = vosk.KaldiRecognizer(model, 16000)
    pa = pyaudio.PyAudio()
    stream = pa.open(format=pyaudio.paInt16, channels=1, rate=16000, input=True, frames_per_buffer=8000)
    stream.start_stream()
    print("Listening...")
    try:
        while True:
            data = stream.read(4000, exception_on_overflow=False)
            if recognizer.AcceptWaveform(data):
                result = json.loads(recognizer.Result())
                text = result.get("text", "")
                if text:
                    answer = answer_question(text)
                    _speak(answer)
    except KeyboardInterrupt:
        pass
    finally:
        stream.stop_stream()
        stream.close()
        pa.terminate()


def main() -> None:
    parser = argparse.ArgumentParser(description="Voice-activated coaching assistant")
    parser.add_argument("--live", action="store_true", help="listen on microphone")
    parser.add_argument("--text", help="text question for quick query")
    parser.add_argument("--context_file", action="append", default=[], help="additional context file")
    args = parser.parse_args()

    if args.live:
        listen_and_respond()
    elif args.text:
        print(answer_question(args.text, args.context_file))
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
